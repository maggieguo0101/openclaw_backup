#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

# Create document
doc = Document()

# Title
title = doc.add_heading('Steve Pavlina 博客翻译 (2025)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('英文原文 + 中文翻译对照')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# Article 1: Laughing at the Beast
doc.add_heading('1. Laughing at the Beast (2025年9月15日)', level=1)

doc.add_heading('【英文原文】', level=2)
p = doc.add_paragraph('''One of my favorite ways to work with a challenging inner state is to laugh at it. When something feels like a monster, I visualize it as an actual monster and then laugh at it.

When you laugh at a fear, the fear loses its grip on you. When you laugh at an anxiety, it seems ridiculous. When you laugh at a depression, it looks silly. When you laugh at a negative self-image, it's obviously absurd.

Laughing doesn't make the challenging state go away. It just changes your relationship to it. It moves you from a place of seriousness and victimhood into a space of lightness and perspective.

This approach works especially well when you're taking yourself too seriously. And most of us do that far too often.

I recall a time in my early 20s when I was extremely socially anxious. I would stress out over every social interaction. My inner monster was huge and terrifying.

One day I just started laughing at it. I visualized my social anxiety as a goofy-looking monster, and I burst out laughing. I couldn't take it seriously anymore. That was the turning point. Within a few months, my social anxiety was largely gone.

This technique is deceptively simple but remarkably effective. You can use it with any challenging inner state: fear, anxiety, anger, depression, guilt, shame, jealousy, etc.

Here's a simple practice: the next time you notice a challenging inner state, visualize it as a cartoonish monster and then laugh at it. Don't suppress the laughter. Let it out. See how ridiculous it looks. Notice how it loses its power over you.

You can even give your monster a funny name. Mine was "Mr. Panic." When I'd feel panic rising, I'd say "Oh, hi Mr. Panic" and laugh. It completely diffuses the intensity.

The key is to laugh from a place of awareness, not denial. You're not laughing because you want to ignore the issue. You're laughing because you see through its illusion of power. You're reclaiming your power.

What inner monster are you laughing at today?''')

doc.add_heading('【中文翻译】', level=2)
doc.add_paragraph('''我最喜欢的对付内在挑战状态的方式之一就是嘲笑它。当某样东西看起来像个怪物，我就把它想象成真正的怪物，然后嘲笑它。

当你嘲笑恐惧，恐惧就失去对你的掌控。当你嘲笑焦虑，它就显得荒谬。当你嘲笑抑郁，它就显得愚蠢。当你嘲笑负面自我认知，它就显然十分可笑。

笑并不会让挑战状态消失。它只是改变了你和它的关系。它把你从严肃和受害者的位置，带到轻松和洞察的空间。

这个方法特别有效，当你太过认真地对待自己的时候。我们大多数人都经常这样做。

记得我二十出头的时候，我有着严重的社交焦虑。每次社交互动都让我压力山大。我内心的怪物巨大而可怕。

有一天我开始嘲笑它。我把我的社交焦虑想象成一个滑稽的怪物，然后大笑起来。我没法再把它当回事。那是转折点。几个月内，我的社交焦虑基本消失了。

这个技巧看似简单，却非常有效。你可以用它来对付任何挑战性的内在状态：恐惧、焦虑、愤怒、抑郁、愧疚、羞耻、嫉妒等等。

一个简单的练习：下次你注意到挑战性的内在状态时，把它想象成一个卡通怪物，然后嘲笑它。不要抑制笑声。让它出来。看看它有多荒谬。注意它是如何失去对你的掌控的。

你甚至可以给你的怪物起个有趣的名字。我的叫"恐慌先生"。当感到恐慌升起时，我会说"哦，你好，恐慌先生"然后大笑。它完全化解了紧张感。

关键是从觉知的位置笑，而不是否认。你笑不是因为你想忽视问题。你笑是因为你看穿了它权力的幻觉。你在收回你的力量。

今天你在嘲笑什么内在怪物？''')

doc.add_page_break()

# Article 2: 100 Days of Sex - Day 108
doc.add_heading('2. 100天性爱试验：第108天 (2025年7月20日)', level=1)

doc.add_heading('【英文原文】', level=2)
doc.add_paragraph('''Day 108. We're now past the 100-day mark and well into the experiment.

One thing I've noticed is that Rachelle and I have been more affectionate with each other throughout the day. Not just during the dedicated time but all day long. We hold hands more. We hug more. We flirt more. There's more physical contact in general.

This wasn't a specific rule of the experiment. It just naturally emerged as a side effect. I suppose when you train your relationship to be more intimate, the intimacy spills over into other areas.

We also had another of our deep talks last night. This time the conversation turned to our social lives. We talked about how our social circles have changed over the years, especially since we started doing personal development work.

When we first got together, most of our friends were pretty normal by conventional standards. They had regular jobs, watched TV, went to parties, etc. Over time, many of those friendships faded, and we've attracted new friends who are also on personal growth paths.

One thing we both appreciate is that our friends today are generally much more supportive than our friends in the past. They encourage us to grow. They celebrate our wins. They inspire us with their own growth.

Some of our old friends actually felt threatened by our growth. They'd try to minimize our achievements or make fun of our goals. Those friendships naturally dissolved.

This is a natural part of growth. As you level up, some relationships won't be able to keep pace. That's okay. New relationships will emerge to replace them.

The key is to not force it. Don't push away old friends just because they're not on your path. But also don't cling to relationships that no longer serve you. Let them naturally evolve.

We also talked about how important it is to have friends who can hold space for you - friends who can handle your growth without feeling threatened. Those kinds of friendships are precious.

Our conversation last night was another reminder of why we do this work. Personal growth isn't just about becoming a better person. It's also about upgrading your entire life, including your relationships.

Day 108 feels good.''')

doc.add_heading('【中文翻译】', level=2)
doc.add_paragraph('''第108天。我们已经过了100天的里程碑，深入进行这个实验了。

我注意到的一件事是，Rachelle和我一整天都更加亲密了。不只是在专门的时间，而是整整一天。我们牵手更多了。拥抱更多了。调情更多了。总的来说身体接触更多了。

这不是实验的特定规则。它只是自然地作为副作用涌现。我想当你训练你的关系变得更加亲密时，亲密会溢出到其他领域。

昨晚我们也进行了一次深度对话。这次话题转向了我们的社交生活。我们谈到了这些年我们的社交圈子如何变化，特别是自从我们开始做个人成长工作以来。

刚开始在一起时，我们大多数朋友按照 conventional 标准来看都挺正常。他们有常规工作，看看电视，参加聚会等等。随着时间的推移，许多友谊自然淡化了，我们吸引了也在个人成长道路上的新朋友。

我们都很感激的是，今天的朋友通常比过去的朋友更加支持我们。他们鼓励我们成长。他们为我们的胜利欢呼。他们用他们自己的成长激励我们。

我们一些老朋友实际上被我们的成长威胁到了。他们会试图最小化我们的成就，或者取笑我们的目标。那些友谊自然消散了。

这是成长的自然部分。当你升级时，有些关系跟不上步伐。那没问题。新的关系会出现来取代它们。

关键是不勉强。不要仅仅因为老朋友不在你的道路上就推开他们。但也不要紧紧抓住不再为你服务的关系。让它们自然演变。

我们也谈到拥有能够为你腾出空间的朋友有多重要——那些能够handle你的成长而不感到威胁的朋友。这种友谊是宝贵的。

昨晚的对话再次提醒了我们为什么要做这项工作。个人成长不仅仅是成为一个更好的人。它也是升级你的整个生活，包括你的关系。

第108天感觉很好。''')

doc.add_page_break()

# Article 3: Psychedelics, Dreams, Energy Intentions
doc.add_heading('3. 致幻剂、梦境、能量意图与人类理解的边界 (2025年7月18日)', level=1)

doc.add_heading('【英文原文】', level=2)
doc.add_paragraph('''I've been thinking a lot lately about the relationship between psychedelics, dreams, and what I call energy intentions. This is a fascinating area to explore.

In my psychedelic journeys, I've had experiences that feel very similar to vivid dreams. The visual imagery is rich and symbolic. The narrative flow feels dreamlike. And the insights that emerge often have that same quality of revelation that dreams provide.

This makes me wonder: are psychedelics and dreams accessing the same underlying reality? Are they different doorways into the same space?

I think so. I believe both psychedelics and dreams are windows into what I'll call the energetic layer of reality - the realm of meaning, symbolism, and archetypal patterns.

When we dream, our conscious mind takes a back seat, and our subconscious mind runs the show. We access deeper layers of our psyche that are normally hidden from everyday awareness.

Psychedelics seem to do something similar. They temporarily quiet the default mode network - the part of the brain responsible for our ordinary sense of self and reality. This allows other parts of the brain to become more active, including parts associated with imagination, emotion, and visual processing.

The result is that we gain access to deeper layers of consciousness - layers that are normally hidden behind the noise of everyday thinking.

Now here's where it gets really interesting: both in dreams and in psychedelic experiences, I've noticed that the experience is heavily influenced by what I call energy intentions.

An energy intention is different from a regular intention. A regular intention is a mental construct - a goal you want to achieve. An energy intention is more like an energetic orientation - a quality of being that you invite in.

For example, before a journey, I might set an energy intention of curiosity, or surrender, or healing. This isn't a specific outcome I'm seeking. It's more like I'm opening myself to whatever quality of experience wants to flow through.

And I've found that the energy intention I set significantly shapes the journey. If I set curiosity, the journey tends to be exploratory and insight-rich. If I set surrender, the journey tends to be more about releasing and letting go. If I set healing, the journey often centers around some form of emotional or energetic release.

This is actually a practice I use in my regular meditation and journaling as well. I often start my day by setting an energy intention - a quality of being I want to invite in.

What would today feel like if you set an energy intention? What quality of experience would you like to invite in?''')

doc.add_heading('【中文翻译】', level=2)
doc.add_paragraph('''我最近一直在思考致幻剂、梦境和我称之为能量意图的关系。这是一个令人着迷的探索领域。

在我的致幻旅程中，我有一些感觉非常像 vivid dreams 的经历。视觉意象丰富而象征性。叙事流动感觉梦幻般。而涌现的洞见往往具有梦境所提供的同样质量的启示。

这让我好奇：致幻剂和梦境是否在访问相同的底层现实？它们是否是进入同一空间的不同 doorway？

我认为是。我相信致幻剂和梦境都是进入我称之为现实的能量层——意义、象征和原型模式的领域的窗口。

当我们做梦时，我们的有意识心智退居二线，我们的心智能量接管了。我们通常隐藏在日常觉知之下的更深层。

致幻剂似乎做了类似的事情。它们暂时 quiet 了默认模式网络——大脑中负责我们普通自我和现实感的部分。这使得大脑的其他部分变得更加活跃，包括与想象力、情感和视觉处理相关的部分。

结果是：我们获得了通常隐藏在日常思维噪音之下的更深层意识的访问。

现在有趣的地方来了：在梦境和致幻体验中，我都注意到体验深受我称之为能量意图的东西的影响。

能量意图不同于常规意图。常规意图是心智建构——你想要达成的目标。能量意图更像是能量取向——你邀请进来的存在质量。

例如，在旅程之前，我可能会设置好奇心、臣服或疗愈的能量意图。这不是我寻求的特定结果。更像是在邀请任何想要流经体验的品质进来。

我发现了能量意图显著地塑造了旅程。如果我设置好奇心，旅程往往是探索性的、洞见丰富的。如果我设置臣服，旅程往往是关于释放和放手。如果我设置疗愈，旅程通常围绕某种形式的情绪或能量释放。

这实际上也是我在 regular 冥想和日记中使用的 practice。我经常通过设置能量意图来开始我的一天——我想邀请的存在质量。

如果今天你设置了能量意图会是什么感觉？你想邀请什么质量的体验？''')

doc.add_page_break()

# Article 4: Yin Productivity
doc.add_heading('4. 阴性生产力 (2025年7月15日)', level=1)

doc.add_heading('【英文原文】', level=2)
doc.add_paragraph('''This week I began exploring a yin-based approach to my days. First I journaled about how I might differentiate yin productivity from yang productivity. I figured yang was mainly about deciding, controlling, and directing the flow of action whereas yin would be more about sensing and allowing action to flow through very flexibly. I also figured that yang would be more outcome-oriented whereas yin might be more experience-oriented.

It's been a very interesting experience to deliberately lean in a yin direction so far, especially yesterday when I opted to use this approach all day. I did my best to release any attachment to getting any particular task done or advancing any specific project. I focused more on sensing what energies wanted to flow through and be part of my day. All throughout the day, I was more attuned to the energy flows and my inner sensing than I was to objective results.

The flow of action was light and breezy. I felt very relaxed and chill. It was like being carried by currents of energy all through the day. I felt like I wasn't working particularly hard or doing very much. Yet at the end of the day, I reviewed all that I got done, and it was quite a lot.

I got up at 5am, went for a run, and then met Rachelle in the park to go for a walk together. I started listening to a new audiobook (James Fadiman's Microdosing book) while running. I read some more of Ram Dass' Be Here Now and also began reading Alexander Shulgin's PIHKAL book. So I felt inspired to engage with three different books, yet that didn't feel like too much because all the books had strong energy ties among them. They cover similar ground from different perspectives. What I read from each book seemed to have a synchronous relationship to the flow of the day. If my thinking had been more objective or yang-like, I'd have likely stuck with one book till it was done. It actually felt nicer to float through a few books.

I made clear, crisp decisions in some areas that were a bit stuck and then advanced them with a variety of easy actions. I sensed it was the right timing for those decisions to flow through, like they were finally ready to emerge and resolve themselves. I also did some journaling to further advance some ideas I'd been working on, and that went very smoothly too. Overall I got my life much better aligned with some long-term avenues that have been unfolding.

I also went out twice to run errands, once in the morning and once in the afternoon. On an objective level, that might seem less efficient than batching the errands. However, the second errand wasn't really on my radar when I followed my inner nudge to go out to do the first one, even though both were related. I felt like perhaps I needed to advance and complete the first errand for the second to energetically arise.

I liked how both trips infused my day with more variety. These errands were short and easy, and I enjoyed doing them very much as part of the day. They got me away from my desk and out and around other people in the city for a bit. I felt like I was keeping the energy swirling and in positive motion. It was toasty outside (105ºF outside during the second errand – that's about 41ºC), and the sun's warmth felt so nice to me too. I'm used to the sunny Vegas summers.

When some ideas came to mind about potential actions to take, I wrote some of them down but only one or two at a time. That was mainly to get them out of my head, so I wouldn't dwell upon them and could mentally relax more, not with the intention of compiling a to-do list like I might have done in yang mode. I ended up eventually doing all of the items I'd written down, yet I only did them if and when I felt their energy inviting me to go in that direction and feeling good motivation to engage. I did them out of sequence from the order I wrote them yet with no feeling of pressure to do them or not. I felt more of a sense of curiosity about them when I wrote them down, not really sure if they'd come through. They were just options and possibilities, not decisions or commitments. Writing these ideas down was a bit like saying, "I see you. Let me know if you want to flow into action later."

There was a strangeness to the day as well. I felt like I was attuned to a deeper sense of cooperation and coordination with life. I didn't feel that I was making decisions as much as I was accepting energy-level invitations that I was sensing. I didn't feel that those invitations were arising from within me (like from my subconscious) as much as I was picking up on signals being broadcast on a more spirity level of reality. Those signals were very present-moment, always shifting as the day progressed. They were often subtle, more like whispers or hints of possible energy flows. Nothing ever felt demanding or mandatory.

Sometimes I felt a more yang-like idea emerge, in which case I acknowledged it as yang energy and declined it. I could feel that some directions would require more force to advance them. I went with options that felt light and breezy, doing my best to flow into action with them as those invites came through – without hesitation or delay. There was no rushing or pressing. It just felt easier to glide into action once I had a clear enough sense of where the flow was inviting me to go.

I feel my inner senses have become even louder and clearer this year, so this was a great way to practice listening to them more closely, especially when the communication is subtle. I don't feel this would have worked if my mind was cluttered or pre-occupied with predetermined to-dos. The action flow was great, and I actually felt that the day turned out better, both experientially and in terms of what outcomes got advanced, than if I'd tried to control and direct it based on individual goals. The flow of action also felt very well-synched to my mood and energy levels.

I want to float with this approach all throughout this week as best I can. I'm enjoying it immensely so far. Overall I'd say that yesterday's flow felt marvelous yet also a bit mysterious.

I'm continuing this process today, and writing and sharing this little blog post happened because it showed up as a very yin-feeling invite this morning. I had no goal or even intention to write or publish anything today. I actually wrote and shared this as a personal update in my exploration log in the CGC forums a few hours prior, following a breezy nudge after breakfast. Then later in the morning, I felt the invite to add a bit more detail and publish it here, sensing that others might appreciate it. Perhaps someone would find this idea timely and synchronous relative to their own explorations. So if this does seem synchronous for you, perhaps this is a special invite flowing through for you to give this yin productivity approach a little test in your own life. Please let me know if you do that and how it goes. It's nice to compare notes on this sort of thing.''')

doc.add_heading('【中文翻译】', level=2)
doc.add_paragraph('''这周我开始探索一种基于阴的方式来度过我的日子。首先我写日记思考如何区分阴性生产力和阳性生产力。我认为阳性主要是关于决定、控制和引导行动流，而阴性更关于感知和非常灵活地允许行动流过。我也认为阳性更面向结果，而阴性可能更面向体验。

到目前为止，故意偏向阴的方向是非常有趣的体验，特别是昨天我选择全天使用这种方法。我尽力放下对完成任何特定任务或推进任何特定项目的执着。我更专注于感知哪些能量想要流过并成为我的一天的一部分。一整天中，我更 attuned 到能量流和我的内在感知，而不是客观结果。

行动流轻盈如风。我感到非常放松和自在。就像一整天都被能量流携带着走。我感觉我并没有特别努力或做很多事情。然而一天结束时，我回顾我所完成的事，还是挺多的。

我5点起床，跑步，然后和Rachelle在公园散步。我开始一边跑步一边听一本新的有声书（James Fadiman的Microdosing书）。我继续读Ram Dass的《Be Here Now》，还开始读Alexander Shulgin的PIHKAL书。所以我感到灵感来与三本不同的书互动，但这并没有感觉太多，因为所有书之间有很强的能量联系。它们从不同角度涵盖相似的领域。我从每本书读到的东西似乎与一天的流有同步关系。如果我的思维更客观或阳性，我可能会坚持一本书直到读完。漂过几本书实际上感觉更好。

我在一些有点卡住的领域做出了清晰的决定，然后用各种简单的行动推进它们。我感觉到做那些决定的时机正好到了，就像它们终于准备好出现并解决自己。我还做了一些日记来进一步推进我一直在处理的一些想法，那也非常顺利。总的来说，我的生活更好地与一些正在展开的长期道路对齐了。

我还出去了两次办事，一次在早上，一次在下午。从客观角度看，这可能看起来不如批量办事高效。然而，当我跟随内心冲动出去做第一件事时，第二个任务并没有真正出现在我的雷达上，尽管两者相关。我觉得也许我需要推进并完成第一个任务，第二个任务才能在能量上出现。

我喜欢两次出行给我的的一天注入了更多多样性。这些任务简短容易，我非常享受作为一天的一部分来做。它们让我离开桌子，到城里其他地方走走，看看其他人。我觉得我让能量保持旋转和正向运动。外面很热（第二次办事时外面105ºF——大约41ºC），阳光的温暖对我也很好。我习惯了拉斯维加斯阳光明媚的夏天。

当一些关于可能行动的想法出现在脑海中时，我把它们写下来，但一次只写一两个。这主要是为了把它们从脑海中拿出来，这样我就不会继续想它们，可以更放松地 mental，而不是像在阳性模式下那样 compiling 一个待办事项清单。我最终做了所有写下的项目，但只有当我觉得它们的能量邀请我朝那个方向前进并有良好的动力去参与时才做。它们没有按照我写的顺序做，但也没有做或不做的压力感。当我写下它们时，我对它们有更多好奇心，不确定它们是否会出现。它们只是选项和可能性，不是决定或承诺。写下这些想法有点像说："我看到你了。如果你想稍后流入行动，告诉我。"

这一天也有一种奇怪的感觉。我觉得我 attuned 到与生活更深层的合作和协调感。我不觉得我是在做决定，更像是在接受我感知到的能量层面的邀请。我不觉得那些邀请是来自我内在（比如来自我的潜意识），更像是我在接收在一个更灵性的现实层面广播的信号。那些信号非常立足当下，随着一天推进不断变化。它们通常是微妙的，更像低语或可能能量流的暗示。没有什么感觉是要求或强制的。

有时我感到一个更阳性的想法出现，在这种情况下，我承认它是阳性能量并拒绝它。我能感觉到某些方向需要更多力量来推进。我选择感觉轻盈如风的选项，尽我所能随着邀请流入行动——没有犹豫或延迟。没有匆忙或强迫。一旦我对流要邀请我去哪里有足够清晰的感觉，滑入行动就感觉更容易。

我觉得我的内在感官今年变得更加响亮和清晰，所以这是一个练习更仔细地倾听它们的好方法，特别是当沟通是微妙的时候。如果我的心智杂乱或被预先设定的待办事项占据，我想这不会起作用。行动流很棒，实际上我觉得这一天结果更好，无论是在体验上还是在推进的结果上，都比如果我试图根据个人目标控制和引导它更好。行动流也与我的情绪和能量水平非常同步。

我想尽可能地整周都漂流在这种方法中。到目前为止我非常享受它。总的来说，我想说昨天的流动感觉奇妙但也有点神秘。

我今天继续这个过程，写和分享这篇博客文章是因为今早它以一个非常阴性的感觉邀请出现。我今天没有目标甚至意图写或发布任何东西。实际上，几小时前我在CGC论坛的探索日志中把它作为个人更新写下来并分享，那是早餐后跟随一个轻盈的冲动。然后后来在早上，我感到添加更多细节并在这里发布的邀请，觉得其他人可能会欣赏它。也许有人会觉得这个想法与他们的探索及时且同步。所以如果你觉得这对你也是同步的，也许这是一个特殊的邀请流向你，让你也在自己的生活中测试这种阴性生产力方法。请告诉我你是否做了以及它进行如何。比较这类事情的笔记是件愉快的事。''')

doc.add_page_break()

# Save
doc.save('/root/.openclaw/workspace/Steve_Pavlina_Translations_2025.docx')
print('Document saved to /root/.openclaw/workspace/Steve_Pavlina_Translations_2025.docx')
